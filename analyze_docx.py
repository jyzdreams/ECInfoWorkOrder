import docx

doc = docx.Document('D:\\ec-workorder\\temp.docx')
table = doc.tables[0]

print(f"=== 文档基本信息 ===")
print(f"标题: {doc.paragraphs[0].text}")
print(f"表格行数: {len(table.rows)}")
print(f"表格列数: {len(table.columns)}")
print(f"\n=== 表头列名 ===")
for i, cell in enumerate(table.rows[0].cells):
    print(f"列{i}: {cell.text}")

print(f"\n=== 完整表格内容 ===")
for row_idx, row in enumerate(table.rows):
    print(f"\n--- 第{row_idx + 1}行 ---")
    for col_idx, cell in enumerate(row.cells):
        text = cell.text.strip()
        if text:
            print(f"  {table.rows[0].cells[col_idx].text}: {text}")
